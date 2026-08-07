import os
import email
import fitz
from docx import Document
from pathlib import Path
from bs4 import BeautifulSoup

def load_docx(file_path: str) -> dict:
    doc = Document(file_path)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                tables_text.append(row_text)
    if tables_text:
        full_text += "\n\n--- 表格 ---\n" + "\n".join(tables_text)

    file_name = Path(file_path).stem
    return {
        "text": full_text,
        "metadata": {
            "source": file_path,
            "title": file_name,
            "file_size": os.path.getsize(file_path),
            "file_type": "docx",
        }
    }

def load_pdf(file_path: str) -> dict:
    doc = fitz.open(file_path)
    pages = []
    for page in doc:
        text = page.get_text()
        if text and text.strip():
            pages.append(text.strip())
    doc.close()

    file_name = Path(file_path).stem
    full_text = "\n\n".join(pages)
    return {
        "text": full_text,
        "metadata": {
            "source": file_path,
            "title": file_name,
            "file_size": os.path.getsize(file_path),
            "file_type": "pdf",
        }
    }

def load_mime_html(file_path: str) -> dict:
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as fh:
        msg = email.message_from_string(fh.read())

    html_text = None
    for part in msg.walk():
        if part.get_content_type() == 'text/html':
            payload = part.get_payload(decode=True)
            if payload:
                html_text = payload.decode('utf-8', errors='ignore')
                break

    if not html_text:
        raise ValueError("No text/html part found in MIME file")

    soup = BeautifulSoup(html_text, 'lxml')
    for tag in soup(['script', 'style', 'meta', 'link', 'head']):
        tag.decompose()
    text = soup.get_text()
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    full_text = '\n'.join(lines)

    file_name = Path(file_path).stem
    return {
        "text": full_text,
        "metadata": {
            "source": file_path,
            "title": file_name,
            "file_size": os.path.getsize(file_path),
            "file_type": "mime-html",
        }
    }

def is_mime_file(file_path: str) -> bool:
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as fh:
            first_line = fh.readline().lower()
        return first_line.startswith('mime-version:')
    except Exception:
        return False

def load_all_docs(folder: str) -> list[dict]:
    docs = []
    for root, dirs, files in os.walk(folder):
        for f in files:
            fpath = os.path.join(root, f)
            if f.startswith('~$'):
                continue
            try:
                if is_mime_file(fpath):
                    docs.append(load_mime_html(fpath))
                elif f.endswith('.docx'):
                    docs.append(load_docx(fpath))
                elif f.endswith('.pdf'):
                    docs.append(load_pdf(fpath))
                elif f.endswith('.doc'):
                    print(f"  ⚠️ 跳过无法解析的 .doc 文件：{f}")
            except Exception as e:
                print(f"  ⚠️ 加载失败 {f}: {e}")
    return docs
